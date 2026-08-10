from pathlib import Path

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from enterprise_knowledge_rag.documents.extractors import (
    ExtractorRegistry,
    SourceValidationError,
)
from enterprise_knowledge_rag.documents.source_models import (
    ExtractedBlockKind,
    IssueSeverity,
    SourceFile,
)


def write_docx_fixture(path: Path) -> Path:
    document = Document()
    document.add_heading("差旅报销流程", level=1)
    document.add_paragraph("员工应在返程后提交报销申请。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "审批人"
    table.cell(0, 1).text = "额度"
    table.cell(1, 0).text = "部门负责人"
    table.cell(1, 1).text = "5000 元"
    document.save(path)
    return path


def write_pdf_fixture(path: Path, page_texts: list[str]) -> Path:
    pdf = canvas.Canvas(str(path))
    for page_number, text in enumerate(page_texts, start=1):
        pdf.drawString(72, 800, "Enterprise Policy")
        pdf.drawString(72, 740, text)
        pdf.drawString(72, 40, f"Page {page_number}")
        pdf.showPage()
    pdf.save()
    return path


def test_docx_extractor_preserves_heading_and_table(tmp_path: Path) -> None:
    path = write_docx_fixture(tmp_path / "expense-process.docx")

    extracted = ExtractorRegistry.default().extract(SourceFile.from_path(path))

    assert any(
        block.kind is ExtractedBlockKind.HEADING
        and block.text == "差旅报销流程"
        for block in extracted.blocks
    )
    assert any(
        block.kind is ExtractedBlockKind.TABLE and "审批人" in block.text
        for block in extracted.blocks
    )


def test_pdf_extractor_records_pages_and_text(tmp_path: Path) -> None:
    path = write_pdf_fixture(
        tmp_path / "leave-policy.pdf",
        ["Employees submit leave requests before absence." for _ in range(2)],
    )

    extracted = ExtractorRegistry.default().extract(SourceFile.from_path(path))

    assert extracted.page_count == 2
    assert {block.page_number for block in extracted.blocks} == {1, 2}
    assert not any(
        issue.severity is IssueSeverity.BLOCKING for issue in extracted.issues
    )


def test_low_text_pdf_is_marked_blocking_for_manual_handling(tmp_path: Path) -> None:
    path = write_pdf_fixture(tmp_path / "scan.pdf", ["x"])

    extracted = ExtractorRegistry.default().extract(SourceFile.from_path(path))

    assert any(
        issue.code == "scanned_or_low_text_pdf"
        and issue.severity is IssueSeverity.BLOCKING
        for issue in extracted.issues
    )


def test_markdown_and_gb18030_text_are_normalized_to_blocks() -> None:
    markdown = SourceFile.from_bytes(
        original_filename="policy.md",
        media_type="text/markdown",
        content="# 采购制度\n\n供应商需要完成准入审核。".encode(),
    )
    text = SourceFile.from_bytes(
        original_filename="notice.txt",
        media_type="text/plain",
        content=("员工出差前需要提交申请。" * 10).encode("gb18030"),
    )

    markdown_result = ExtractorRegistry.default().extract(markdown)
    text_result = ExtractorRegistry.default().extract(text)

    assert markdown_result.blocks[0].kind is ExtractedBlockKind.HEADING
    assert "供应商" in markdown_result.blocks[1].text
    assert "员工出差前" in text_result.blocks[0].text


def test_extension_mime_and_signature_must_agree() -> None:
    disguised_pdf = SourceFile.from_bytes(
        original_filename="policy.pdf",
        media_type="application/pdf",
        content=b"this is not a PDF",
    )

    with pytest.raises(SourceValidationError, match="signature"):
        ExtractorRegistry.default().extract(disguised_pdf)


def test_source_larger_than_fifteen_mib_is_rejected_before_parsing() -> None:
    oversized = SourceFile.from_bytes(
        original_filename="policy.txt",
        media_type="text/plain",
        content=b"a" * (15 * 1024 * 1024 + 1),
    )

    with pytest.raises(SourceValidationError, match="15 MiB"):
        ExtractorRegistry.default().extract(oversized)
